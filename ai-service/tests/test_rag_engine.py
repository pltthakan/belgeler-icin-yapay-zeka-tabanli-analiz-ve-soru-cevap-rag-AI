import io
import json
import os
import tempfile
import unittest
from unittest.mock import patch

from docx import Document

from rag_engine import RagEngine


class RagEngineAnswerTests(unittest.TestCase):
    def setUp(self):
        self.previous_ollama_base_url = os.environ.get("OLLAMA_BASE_URL")
        self.previous_ollama_model = os.environ.get("OLLAMA_MODEL")
        os.environ["OLLAMA_BASE_URL"] = ""
        os.environ["OLLAMA_MODEL"] = ""
        self.data_dir = tempfile.TemporaryDirectory()
        self.engine = RagEngine(self.data_dir.name)

    def tearDown(self):
        self.data_dir.cleanup()
        self._restore_environment_variable("OLLAMA_BASE_URL", self.previous_ollama_base_url)
        self._restore_environment_variable("OLLAMA_MODEL", self.previous_ollama_model)

    def _restore_environment_variable(self, name, value):
        if value is None:
            os.environ.pop(name, None)
        else:
            os.environ[name] = value

    def test_document_overview_questions_use_the_same_document_profile(self):
        chunks = [{
            "text": (
                "İNÖNÜ ÜNİVERSİTESİ MÜHENDİSLİK FAKÜLTESİ BİLGİSAYAR MÜHENDİSLİĞİ BÖLÜMÜ\n"
                "İŞLETMEDE MESLEKİ EĞİTİM DERSİ\n"
                "ÖĞRENCİLERİN İŞLETME DEĞERLENDİRME ANKETİ\n"
                "20- Eğitim bittikten sonra çalışmak ister misiniz?"
            )
        }]
        sources = [{"text": "20- Eğitim bittikten sonra çalışmak ister misiniz?"}]
        document_profile = self.engine._build_document_profile(chunks)

        answers = [
            self.engine._build_answer(question, sources, document_profile)
            for question in (
                "Bu belgenin ana konusu nedir?",
                "Bu belge nedir?",
                "Bu belgede ne anlatılıyor?",
                "Belgeyi özetle.",
            )
        ]

        expected_answer = (
            "Bu belge, “İŞLETMEDE MESLEKİ EĞİTİM DERSİ — "
            "ÖĞRENCİLERİN İŞLETME DEĞERLENDİRME ANKETİ” başlıklı bir ankettir."
        )
        self.assertEqual(answers, [expected_answer] * 4)

    def test_topic_question_detects_turkish_dotted_and_dotless_i(self):
        self.assertTrue(self.engine._is_document_overview_question("Bu belgenin ana konusu nedir?"))
        self.assertTrue(self.engine._is_document_overview_question("Belge ne hakkında?"))
        self.assertTrue(self.engine._is_document_overview_question("Bu belgede ne anlatılıyor?"))
        self.assertTrue(self.engine._is_document_overview_question("Bu nasıl bir belge?"))
        self.assertTrue(self.engine._is_document_overview_question("Belgenin içeriği ne?"))
        self.assertTrue(self.engine._is_document_overview_question("Bu belgede neler var?"))

    def test_response_mode_distinguishes_summary_critique_and_factual_questions(self):
        self.assertEqual(self.engine._classify_response_mode("Bu belgeyi özetle."), "summary")
        self.assertEqual(self.engine._classify_response_mode("Bu CV'nin eksikleri neler?"), "critique")
        self.assertEqual(self.engine._classify_response_mode("Ödeme süresi kaç gün?"), "factual")

    def test_critique_fallback_does_not_present_an_inference_as_a_fact(self):
        answer = self.engine._build_answer(
            "Bu CV'nin eksikleri neler?",
            [{"text": "Adayın Java ve React projeleri bulunuyor."}],
            {"title": "Aday CV'si", "summary": "Adayın teknik projelerini içeren CV."},
        )

        self.assertIn("yerel LLM", answer)
        self.assertNotIn("negatif taraf", answer.lower())

    def test_extractive_fallback_returns_relevant_sentences_only(self):
        passage = self.engine._extract_relevant_passage(
            "Ödenen ücrette memnun musunuz?",
            "Program sonunda çalışmak ister misiniz? Ödenen ücrette memnun musunuz? EVET HAYIR.",
        )

        self.assertEqual(passage, "Ödenen ücrette memnun musunuz?")

    def test_latest_course_question_uses_transcript_term_order(self):
        chunks = [
            {
                "chunkIndex": 0,
                "pageNumber": 1,
                "score": 1.0,
                "text": (
                    "Ders Kodu Ders Adı Krd2024-2025 Güz\n"
                    "BİLM289-2022 Veri Organizasyonu 3 B1 16,25 5\n"
                    "Ders Kodu Ders Adı Krd2025-2026 Bahar\n"
                    "BİLM468-2022 İşletmede Mesleki Eğitim 5 A1 120 30"
                ),
            }
        ]

        result = self.engine._answer_latest_course_question("Aldığı son ders nedir?", chunks, 0)

        self.assertIsNotNone(result)
        self.assertIn("2025-2026 Bahar", result["answer"])
        self.assertIn("BİLM468-2022 İşletmede Mesleki Eğitim", result["answer"])
        self.assertEqual(result["trace"]["provider"], "transcript-structure")

    def test_order_sensitive_questions_use_document_order_for_any_document_type(self):
        chunks = [
            {"chunkIndex": 0, "text": "İlk bölüm: başvuru koşulları."},
            {"chunkIndex": 1, "text": "Orta bölüm: değerlendirme süreci."},
            {"chunkIndex": 2, "text": "Son bölüm: kararın tebliği."},
        ]

        self.assertEqual(self.engine._order_sensitive_direction("Son bölüm nedir?"), "last")
        self.assertEqual(self.engine._order_sensitive_direction("İlk bölüm nedir?"), "first")
        self.assertEqual(
            [source["chunkIndex"] for source in self.engine._ordered_sources(chunks, "last", 2)],
            [1, 2],
        )
        self.assertEqual(
            [source["chunkIndex"] for source in self.engine._ordered_sources(chunks, "first", 2)],
            [0, 1],
        )

    def test_date_query_expansion_does_not_trigger_document_order_answer(self):
        document_id = "contract-notice-period"
        chunks = [
            {
                "chunkIndex": 0,
                "pageNumber": 1,
                "text": "SÖZLEŞME ÖZETİ\nTaraflar hizmet kapsamı ve ödeme koşullarında mutabık kalmıştır.",
            },
            {
                "chunkIndex": 1,
                "pageNumber": 2,
                "text": (
                    "FESİH\nTaraflardan biri sözleşmeyi feshetmek isterse diğer tarafa "
                    "en az 30 gün önce yazılı bildirim yapmalıdır."
                ),
            },
        ]
        embeddings = self.engine._hashing_vectorizer.transform([chunk["text"] for chunk in chunks]).toarray()
        self.engine._embed_texts = lambda texts: self.engine._hashing_vectorizer.transform(texts).toarray()
        self.engine.disable_qa_model = True
        self.engine._index_path(document_id).write_text(
            json.dumps({
                "documentId": document_id,
                "filename": "contract.txt",
                "chunkCount": len(chunks),
                "chunks": chunks,
                "embeddings": embeddings.tolist(),
                "documentProfile": self.engine._build_document_profile(chunks),
            }, ensure_ascii=False),
            encoding="utf-8",
        )

        result = self.engine.answer_question(
            document_id=document_id,
            question="Sözleşmenin fesih bildirim süresi kaç gündür?",
            top_k=3,
        )

        self.assertEqual(result["sources"][0]["chunkIndex"], 1)
        self.assertIn("30 gün", result["answer"])

    def test_memory_hybrid_search_combines_dense_and_sparse_matches(self):
        chunks = [
            {"chunkIndex": 0, "pageNumber": 1, "text": "Bu bölüm genel eğitim bilgilerini açıklar."},
            {"chunkIndex": 1, "pageNumber": 1, "text": "Aday Pluton projesinde REST API geliştirdi."},
            {"chunkIndex": 2, "pageNumber": 1, "text": "Bu bölüm iletişim bilgilerini içerir."},
        ]
        dense_scores = [0.92, 0.31, 0.12]

        sources = self.engine._hybrid_sources_from_memory(
            chunks=chunks,
            dense_scores=dense_scores,
            question="Pluton REST API",
            top_k=2,
        )

        self.assertEqual(sources[0]["chunkIndex"], 1)
        self.assertEqual(sources[0]["retrievalStrategy"], "hybrid")
        self.assertGreater(sources[0]["sparseScore"], 0)
        self.assertIn(0, [source["chunkIndex"] for source in sources])

    def test_reranker_reorders_retrieval_candidates(self):
        class FakeReranker:
            def predict(self, pairs, show_progress_bar=False):
                return [0.1, 0.9, 0.2]

        self.engine.reranker_enabled = True
        self.engine._reranker_model = FakeReranker()
        sources = [
            {"chunkIndex": 0, "text": "Ücret ve ödeme zamanı."},
            {"chunkIndex": 1, "text": "Rekabet etmeme ve gizlilik hükümleri."},
            {"chunkIndex": 2, "text": "Çalışma süreleri."},
        ]

        reranked = self.engine._rerank_sources(
            "işten ayrıldıktan sonra aynı proje",
            sources,
            top_k=2,
        )

        self.assertEqual([source["chunkIndex"] for source in reranked], [1, 2])
        self.assertTrue(reranked[0]["reranked"])
        self.assertEqual(reranked[0]["rerankerScore"], 0.9)

    def test_lexical_retrieval_score_tolerates_turkish_suffixes(self):
        score = self.engine._lexical_retrieval_score(
            {"ortalama"},
            "Hakan Polat'ın 2022-2023 Bahar dönemi not ortalaması 0,70.",
        )

        self.assertGreater(score, 0)

    def test_qa_quality_gate_rejects_single_character_spans(self):
        self.assertFalse(self.engine._is_usable_qa_answer("I"))
        self.assertFalse(self.engine._is_usable_qa_answer(" "))
        self.assertTrue(self.engine._is_usable_qa_answer("30"))
        self.assertTrue(self.engine._is_usable_qa_answer("Öğrencilerin hakları"))

    def test_generated_answer_sanitizer_removes_answer_and_source_labels(self):
        answer = self.engine._sanitize_generated_answer(
            "Cevap: Adayın backend deneyimi öne çıkıyor. Kaynak 1 ve 2'de projeler yer alıyor."
        )

        self.assertEqual(answer, "Adayın backend deneyimi öne çıkıyor.")

    def test_generated_answer_sanitizer_repairs_pdf_unicode_artifacts(self):
        answer = self.engine._sanitize_generated_answer(
            "Genel Not Ortalaması + Üniversiteye Giri/uni015F Sırası teknik mülakat ba/uni015Farı puanını belirler."
        )

        self.assertEqual(
            answer,
            "Genel Not Ortalaması + Üniversiteye Giriş Sırası teknik mülakat başarı puanını belirler.",
        )

    def test_gano_question_expands_to_general_grade_average_terms(self):
        retrieval_question = self.engine._normalize_question_for_retrieval("gano önemlimi başvuruda")

        self.assertIn("genel not ortalamasi", retrieval_question)
        self.assertIn("universiteye giris sirasi", retrieval_question)

    def test_critique_guard_rejects_claims_contradicted_by_sources(self):
        sources = [{
            "text": (
                "EXPERIENCE\nSoftware Developer Intern\nFeb 2026 – Jun 2026\n"
                "SKILLS\nJava, Spring Boot, Python, Flask, React\n"
                "PROJECTS\nBuilt a financial dashboard and implemented REST APIs."
            )
        }]
        unsupported = (
            "Belgeye dayalı değerlendirme: deneyimlerin tarihlendirilmemiş olması, "
            "teknik becerilerin eksik belirtilmesi ve proje açıklamalarının tamamlanmamış olması."
        )

        self.assertFalse(self.engine._is_grounded_critique(unsupported, sources))
        self.assertTrue(self.engine._is_grounded_critique(self.engine._safe_critique_answer(), sources))

    def test_grounding_guard_accepts_terms_joined_during_pdf_extraction(self):
        sources = [{
            "text": (
                "Kimlik doğrulamaOIDC/OAuth2, SSO, MFA. "
                "YetkilendirmeRBAC + ABAC. "
                "ŞifrelemeTLS everywhere, at-rest encryption."
            )
        }]
        answer = (
            "Güvenlik önlemleri arasında OIDC/OAuth2, SSO, MFA, "
            "RBAC + ABAC yetkilendirme ve TLS şifreleme kullanılmıştır."
        )

        self.assertTrue(
            self.engine._is_grounded_answer(
                "güvenlik önlemleri nelerdir?",
                answer,
                sources,
                "factual",
            )
        )

    def test_evidence_guard_accepts_exact_percentage_entity_and_relation_match(self):
        sources = [{
            "chunkIndex": 12,
            "text": (
                "YKS indirimli kontenjanlar: İlgili yılın YKS kılavuzunda burslu ve ücretlinin "
                "dışında, %50 veya %25 indirimli kontenjanlardır."
            ),
        }]

        decision = self.engine._evidence_support_decision(
            "YKS'de %50 indirimli kontenjana yerleşen öğrenci hangi indirimi alır?",
            sources,
        )

        self.assertTrue(decision["supported"])
        self.assertEqual(decision["reason"], "exact-percentage-entity-relation-match")
        self.assertEqual(decision["chunkIndex"], 12)
        self.assertEqual(decision["matchedPercentages"], ["50"])
        self.assertIn("yks", decision["matchedEntities"])
        self.assertIn("discount", decision["matchedRelations"])

    def test_evidence_guard_rejects_same_percentage_for_unrelated_entity(self):
        sources = [{
            "chunkIndex": 20,
            "text": "Engelli indirimi: Engel oranı %40 ve üzerinde olan öğrencilere %50 indirim uygulanır.",
        }]

        decision = self.engine._evidence_support_decision(
            "YKS'de %50 indirimli kontenjana yerleşen öğrenci hangi indirimi alır?",
            sources,
        )

        self.assertFalse(decision["supported"])
        self.assertEqual(decision["reason"], "percentage-without-entity-relation-support")

    def test_model_no_answer_is_replaced_when_retrieved_evidence_supports_claim(self):
        question = "YKS'de %50 indirimli kontenjana yerleşen öğrenci hangi indirimi alır?"
        sources = [{
            "chunkIndex": 12,
            "text": (
                "YKS indirimli kontenjanlar: İlgili yılın YKS kılavuzunda burslu ve ücretlinin "
                "dışında, %50 veya %25 indirimli kontenjanlardır."
            ),
        }]
        generation = {
            "provider": "ollama",
            "model": "qwen3:8b",
            "responseMode": "factual",
            "prompt": "test",
        }

        with patch.object(
            self.engine,
            "_answer_with_ollama",
            return_value=("Bu bilgi belgede yer almıyor.", generation),
        ):
            with self.assertLogs("uvicorn.error.rag.guard", level="INFO") as logs:
                answer, result_generation = self.engine._build_answer_result(question, sources, {})

        self.assertEqual(
            answer,
            "Belgeye göre YKS'de %50 indirimli kontenjana yerleşen öğrenci %50 öğrenim ücreti indirimi alır.",
        )
        self.assertEqual(result_generation["provider"], "evidence-supported-fallback")
        self.assertEqual(
            result_generation["guardReason"],
            "model-no-answer-despite-supported-evidence",
        )
        self.assertIn("exact-percentage-entity-relation-match", str(result_generation["evidenceDecision"]))
        self.assertIn("model-no-answer-despite-supported-evidence", " ".join(logs.output))

    def test_no_answer_is_allowed_only_without_supporting_evidence(self):
        question = "YKS'de %50 indirimli kontenjana yerleşen öğrenci hangi indirimi alır?"
        sources = [{"chunkIndex": 1, "text": "Yatay geçiş başvuruları akademik takvimde ilan edilir."}]
        generation = {
            "provider": "ollama",
            "model": "qwen3:8b",
            "responseMode": "factual",
            "prompt": "test",
        }

        with patch.object(
            self.engine,
            "_answer_with_ollama",
            return_value=("Bu bilgi belgede yer almıyor.", generation),
        ):
            answer, result_generation = self.engine._build_answer_result(question, sources, {})

        self.assertEqual(answer, "Bu bilgi belgede yer almıyor.")
        self.assertEqual(
            result_generation["guardReason"],
            "no-retrieved-chunk-supports-main-claim",
        )

    def test_docx_extraction_includes_table_cells_in_document_order(self):
        document = Document()
        document.add_paragraph("Başlık")
        table = document.add_table(rows=1, cols=2)
        table.cell(0, 0).text = "Ödeme süresi"
        table.cell(0, 1).text = "30 gün"
        document.add_paragraph("Son not")
        raw_bytes = io.BytesIO()
        document.save(raw_bytes)

        pages = self.engine._extract_docx_pages(raw_bytes.getvalue())

        self.assertEqual(pages[0]["text"], "Başlık\nÖdeme süresi | 30 gün\nSon not")

    def test_semantic_chunking_keeps_paragraph_boundaries(self):
        pages = [{
            "pageNumber": 1,
            "text": (
                "GİRİŞ\n"
                "Bu bölüm sistemin amacını ve kapsamını açıklar. "
                "Kullanıcıların belge yükleyip soru sorabildiğini belirtir.\n\n"
                "YÖNTEM\n"
                "Bu bölüm metin çıkarma, bölümleme ve arama adımlarını açıklar. "
                "Kaynak parçalarının anlam bütünlüğü korunur.\n\n"
                "SONUÇ\n"
                "Bu bölüm sistemin belgeye dayalı cevap ürettiğini özetler. "
                "Yanıtların kaynaklarla birlikte gösterildiğini ve alakasız soruların elendiğini belirtir."
            ),
        }]

        chunks = self.engine._chunk_pages(pages, chunk_size=170, overlap=0)

        self.assertEqual(len(chunks), 3)
        self.assertTrue(chunks[0]["text"].startswith("GİRİŞ"))
        self.assertTrue(chunks[1]["text"].startswith("YÖNTEM"))
        self.assertTrue(chunks[2]["text"].startswith("SONUÇ"))
        self.assertNotIn("YÖ", chunks[0]["text"][-10:])

    def test_heading_aware_chunking_keeps_sections_together(self):
        pages = [{
            "pageNumber": 1,
            "text": (
                "1. Amaç\n"
                "Bu yönergenin amacı eğitim sürecini ve belge işleme kurallarını açıklamaktır.\n"
                "Amaç bölümü ayrı bir kaynak parçası olarak kalmalıdır.\n"
                "2. Kapsam\n"
                "Bu bölüm öğrencileri, eğitmenleri ve eğitim materyallerini kapsar.\n"
                "Kapsam bilgisi aynı başlık altında tutulmalıdır.\n"
                "3. Eğitmenler\n"
                "Eğitmenler Dr. Hüseyin ARIK, Fehmi ARIK ve Erdem Taha Sokullu'dur.\n"
                "Bu kişilerin isimleri başlıktan kopmamalıdır."
            ),
        }]

        chunks = self.engine._chunk_pages(pages, chunk_size=260, overlap=80)

        self.assertEqual(len(chunks), 3)
        self.assertTrue(chunks[0]["text"].startswith("1. Amaç"))
        self.assertTrue(chunks[1]["text"].startswith("2. Kapsam"))
        self.assertTrue(chunks[2]["text"].startswith("3. Eğitmenler"))
        self.assertIn("Dr. Hüseyin ARIK", chunks[2]["text"])

    def test_heading_aware_chunking_falls_back_when_heading_structure_is_weak(self):
        pages = [{
            "pageNumber": 1,
            "text": (
                "Bu belge başlık yapısı olmayan normal bir açıklama metnidir. "
                "Metin birkaç cümleden oluşur ve var olan semantik chunking akışıyla bölünmelidir. "
                "Bu fallback davranışı başlık tespiti zayıf olduğunda korunmalıdır."
            ),
        }]

        self.assertEqual(self.engine._chunk_pages_by_headings(pages, chunk_size=120), [])
        chunks = self.engine._chunk_pages(pages, chunk_size=120, overlap=0)

        self.assertGreaterEqual(len(chunks), 1)

    def test_semantic_chunking_splits_large_blocks_on_word_boundaries(self):
        long_paragraph = (
            "Bu uzun bölüm birinci cümlede sistemin amacını detaylı biçimde anlatır. "
            "İkinci cümle belge parçalarının cümle veya kelime sınırında bölünmesini bekler. "
            "Üçüncü cümle karakter ortasında kesilmiş anlamsız parçalar oluşmamasını doğrular."
        )

        pieces = self.engine._split_oversized_block(long_paragraph, chunk_size=95)

        self.assertGreater(len(pieces), 1)
        self.assertTrue(all(len(piece) <= 95 for piece in pieces))
        self.assertTrue(all(piece == piece.strip() for piece in pieces))
        self.assertFalse(any(piece.endswith(" biç") for piece in pieces))

    def test_pdf_page_extraction_falls_back_when_pypdf_fails(self):
        class BrokenPage:
            def extract_text(self):
                raise ValueError("Odd-length string")

        with patch.object(
            self.engine,
            "_extract_pdf_page_text_with_pymupdf",
            return_value="Fallback PDF metni",
        ) as fallback:
            text = self.engine._extract_pdf_page_text(BrokenPage(), b"%PDF", 1)

        self.assertEqual(text, "Fallback PDF metni")
        fallback.assert_called_once_with(b"%PDF", 1)

    def test_relevance_guard_rejects_low_score_without_calling_a_model(self):
        result = self.engine._relevance_guard_result(
            "Sözleşme hangi şehirde imzalandı?",
            [{"score": 0.04, "text": "Öğrencinin ders notları listelenmiştir."}],
        )

        self.assertIsNotNone(result)
        self.assertEqual(result["provider"], "retrieval-guard")
        self.assertEqual(result["guardReason"], "low-retrieval-score")

    def test_relevance_guard_rejects_obvious_gibberish_even_with_an_accidental_match(self):
        result = self.engine._relevance_guard_result(
            "sjdhfsahjd",
            [{"score": 0.72, "text": "Öğrencinin ders notları listelenmiştir."}],
        )

        self.assertIsNotNone(result)
        self.assertEqual(result["guardReason"], "gibberish-question")

    def test_relevance_guard_rejects_low_information_chat_noise(self):
        sources = [{"score": 0.72, "text": "Hakan Polat'ın 2022-2023 Bahar dönemi not ortalaması 0,70."}]

        for question in ("la get", "ne bileyim ben"):
            with self.subTest(question=question):
                result = self.engine._relevance_guard_result(question, sources)

                self.assertIsNotNone(result)
                self.assertEqual(result["guardReason"], "low-information-question")

    def test_relevance_guard_keeps_short_queries_that_match_document_terms(self):
        result = self.engine._relevance_guard_result(
            "ortalama",
            [{"score": 0.72, "text": "Hakan Polat'ın 2022-2023 Bahar dönemi not ortalaması 0,70."}],
        )

        self.assertIsNone(result)

    def test_relevance_guard_keeps_route_questions_without_exact_term_overlap(self):
        result = self.engine._relevance_guard_result(
            "nerden nereye uçuyor",
            [{"score": 0.72, "text": "MALATYA/MLX\n\nISTANBUL/IST\n\nTK\n2631"}],
        )

        self.assertIsNone(result)

    def test_relevance_guard_keeps_document_overview_questions(self):
        for question in (
            "Bu belgenin ana konusu nedir?",
            "Bu nasıl bir belge?",
            "Belgenin içeriği ne?",
            "Bu belgede neler var?",
        ):
            with self.subTest(question=question):
                result = self.engine._relevance_guard_result(
                    question,
                    [{"score": 0.01, "text": "Öğrencinin ders notları listelenmiştir."}],
                )

                self.assertIsNone(result)


if __name__ == "__main__":
    unittest.main()
